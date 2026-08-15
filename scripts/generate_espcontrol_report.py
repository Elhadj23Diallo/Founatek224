import os
from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
APP = ROOT / 'espcontrol'
OUTPUT = ROOT / 'espcontrol_report.docx'

FILES_TO_INCLUDE = [
    APP / 'models.py',
    APP / 'views.py',
    APP / 'serializers.py',
    APP / 'tests.py',
    APP / 'templates' / 'espcontrol' / 'dashboardUniv.html',
    APP / 'templates' / 'espcontrol' / 'home.html',
]

def read_file_safe(p: Path):
    try:
        return p.read_text(encoding='utf-8')
    except Exception as e:
        return f'Could not read {p}: {e}'

def add_code_block(doc: Document, title: str, content: str):
    doc.add_heading(title, level=2)
    p = doc.add_paragraph()
    run = p.add_run(content)
    font = run.font
    font.name = 'Consolas'
    font.size = Pt(9)

def main():
    doc = Document()
    doc.add_heading('Rapport d\u2019audit - application espcontrol', level=1)
    doc.add_paragraph('Date:')
    doc.add_paragraph(str(Path().resolve()))

    # High-level summary
    doc.add_heading('Résumé', level=2)
    doc.add_paragraph('Ce document décrit le fonctionnement de l\'application Django `espcontrol`.')

    # Include key files
    doc.add_heading('Fichiers clés (extraits)', level=2)
    for p in FILES_TO_INCLUDE:
        if p.exists():
            content = read_file_safe(p)
            # truncate huge files for doc readability
            excerpt = content[:3000]
            if len(content) > 3000:
                excerpt += '\n\n... (fichier tronqué pour le rapport) ...'
            add_code_block(doc, p.relative_to(ROOT).as_posix(), excerpt)
        else:
            doc.add_paragraph(f'Fichier manquant: {p}')

    # Endpoints summary
    doc.add_heading('Endpoints importants', level=2)
    endpoints = [
        ('/api/iot/ingest/', 'Ingestion des données IoT (SensorIngestAPIView)'),
        ('/api/device/<id>/last10/', 'DeviceLast10APIView - renvoie 10 dernières mesures (ISO datetimes, is_anomaly)'),
        ('/api/alerts/latest/', 'latest_alerts - JSON des alertes non-lues'),
        ('/api/alerts/graph/<sensor>/', 'alert_graph_data - historique d\'alertes par capteur'),
        ('/tableau_bord_univ/', 'dashboard_univers - vue HTML du dashboard (protégé)'),
    ]
    for path, desc in endpoints:
        doc.add_paragraph(f'- {path}: {desc}')

    # Agent and background tasks
    doc.add_heading('Agent & traitement asynchrone', level=2)
    doc.add_paragraph('La logique de l\'agent (observer/rules/actions) se trouve dans le package `espcontrol.agent`.')
    doc.add_paragraph('Celery est scaffoldé mais non obligatoire en local; tasks ont un fallback synchrone.')

    # Tests
    doc.add_heading('Tests', level=2)
    doc.add_paragraph('Unité(s) présentes dans `espcontrol/tests.py` couvrant ingestion et DeviceLast10APIView. Commande pour exécuter les tests:')
    doc.add_paragraph('python manage.py test espcontrol -v2')

    # Runbook
    doc.add_heading('Runbook / Mise en route', level=2)
    doc.add_paragraph('1) Créer un virtualenv, installer requirements.txt')
    doc.add_paragraph('2) Lancer migrations: python manage.py migrate')
    doc.add_paragraph('3) Lancer le serveur: python manage.py runserver')

    # Recommendations
    doc.add_heading('Recommandations', level=2)
    doc.add_paragraph('- Remplacer SQLite par Postgres en production')
    doc.add_paragraph('- Déployer un broker Redis pour Celery et workers dédiés')
    doc.add_paragraph("- Gérer l\'état 'is_read' côté client/API pour éviter marquage prématuré")
    doc.add_paragraph('- Ajouter tests d\'intégration pour le dashboard (Selenium/Playwright)')

    doc.add_heading('Annexes', level=2)
    doc.add_paragraph('Liste des fichiers inclus et leur chemin relatifs.')
    for p in FILES_TO_INCLUDE:
        doc.add_paragraph('-', style='List Bullet')
        doc.add_paragraph(p.relative_to(ROOT).as_posix())

    doc.save(OUTPUT)
    print('Report written to', OUTPUT)

if __name__ == '__main__':
    main()
